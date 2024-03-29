from typing import Optional, List, Sequence, Union

import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor

from .base import BaseOutput
from .styles import SimpleTextStyle, FormattedTextStyle
from .text_styles import FormattedText
from .tmp_file import TempOutputFile


class SimpleTextBlock(BaseOutput):
    def __init__(self, text: str, style: Optional[SimpleTextStyle] = None):
        self.text = text
        self.style = style if style else SimpleTextStyle()

    def render_pdf(self):
        pass

    def render_docx(self, document):
        paragraph = document.add_paragraph()

        if self.style.alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif self.style.alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif self.style.alignment == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif self.style.alignment == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            raise ValueError(
                'Unknown alignment parameter.'
            )

        run = paragraph.add_run(self.text)
        run.font.size = Pt(self.style.size)

        if self.style.font_style is None:
            pass
        elif self.style.font_style == 'bold':
            run.font.bold = True
        elif self.style.font_style == 'italic':
            run.font.italic = True
        else:
            raise ValueError(
                'Unknown font style.'
            )


class TokenizedTextBlock(BaseOutput):
    def __init__(self,
                 tokens: List[str],
                 alignment: Optional[str] = 'left',
                 font_size: Optional[int] = 12,
                 highlights: Optional[dict] = None,
                 colors: Optional[dict] = None,
                 ):

        self.tokens = tokens
        self.alignment = alignment
        self.font_size = font_size
        self.highlights = highlights
        self.colors = colors

    def render_pdf(self):
        pass

    def render_docx(self, document):
        paragraph = document.add_paragraph()

        if self.alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif self.alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif self.alignment == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif self.alignment == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            raise ValueError(
                'Unknown alignment parameter.'
            )

        for token in self.tokens:
            run = paragraph.add_run(token)
            run.font.size = Pt(self.font_size)

            if self.highlights:
                if token in self.highlights:
                    hc = WD_COLOR_INDEX.TURQUOISE if self.highlights[token] > 0.0 else WD_COLOR_INDEX.RED
                    run.font.highlight_color = hc


class FormattedTextBlock(BaseOutput):
    def __init__(self, elements: List[FormattedText], style: Optional[FormattedTextStyle] = None):
        self.elements = elements
        self.style = style if style else FormattedTextStyle()

    def render_pdf(self):
        pass

    def render_docx(self, document):
        paragraph = document.add_paragraph()

        if self.style.alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif self.style.alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif self.style.alignment == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif self.style.alignment == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            raise ValueError(
                'Unknown alignment parameter.'
            )

        for run_obj in self.elements:
            run_obj.render_docx(paragraph)


class DescriptiveTextBlock(BaseOutput):
    def __init__(self, text: str, alignment: str = 'left', fontsize: int = 12):
        self.text = text
        self.alignment = alignment
        self.fontsize = fontsize

    def render_pdf(self):
        pass

    def render_docx(self, document):
        paragraph = document.add_paragraph()

        if self.alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif self.alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif self.alignment == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif self.alignment == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            raise ValueError(
                'Unknown alignment parameter.'
            )

        run = paragraph.add_run(self.text)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(self.fontsize)


class SimpleImage(BaseOutput):
    def __init__(self, source: Union[str, TempOutputFile], width=None):
        self.source = source
        self.width = width

    def render_pdf(self):
        pass

    def render_docx(self, document):
        self.width = Inches(self.width) if self.width is not None else self.width

        if isinstance(self.source, str):
            document.add_picture(self.source, width=self.width)

        elif isinstance(self.source, TempOutputFile):
            document.add_picture(self.source.get_path(), width=self.width)
            self.source.delete_file()
        else:
            raise ValueError(
                "Incorrect image source. Image sources must be a valid file path or a TempOutputFile object.")


class Table(BaseOutput):
    def __init__(self,
                 header: List[str],
                 records: List[Sequence],
                 style: str = 'Table Grid',
                 header_fontsize=10,
                 cell_fontsize=9
                 ):

        if records:
            for rec in records:
                if not len(header) == len(rec):
                    raise ValueError(
                        'dimension mismatch between table header and records.'
                    )
        self.header = header
        self.records = records
        self.style = style
        self.header_fontsize = header_fontsize
        self.cell_fontsize = cell_fontsize

    def render_pdf(self):
        pass

    def render_docx(self, document):
        table = document.add_table(rows=len(self.records)+1, cols=len(self.header), style=self.style)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(self.header):
            hdr_cells[i].text = col_name
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(self.header_fontsize)

        for i, rec in enumerate(self.records):
            row_cells = table.rows[i+1].cells
            for j in range(len(self.header)):
                s = f'{rec[j]:.2f}' if isinstance(rec[j], float) else str(rec[j])
                row_cells[j].text = s
                row_cells[j].paragraphs[0].runs[0].font.size = Pt(self.cell_fontsize)


class ImageTable(BaseOutput):
    def __init__(self,
                 images: List,
                 titles: Optional[List[str]] = None,
                 n_cols: int = 2,
                 width: Optional[float] = None,
                 title_fontsize: Optional[int] = 14,
                 ):
        self.images = images
        self.titles = titles
        self.n_cols = n_cols
        self.width = width
        self.title_fontsize = title_fontsize

    def render_pdf(self):
        pass

    def render_docx(self, document):
        self.width = Inches(self.width) if self.width is not None else self.width

        if self.titles:
            if not len(self.titles) == len(self.images):
                raise ValueError("The number of titles must match the number of images.")

        n_rows = int(np.ceil(len(self.images) / self.n_cols))

        table = document.add_table(rows=n_rows, cols=self.n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for img_idx, img in enumerate(self.images):
            row_idx = img_idx // self.n_cols
            col_idx = img_idx % self.n_cols
            paragraph = table.rows[row_idx].cells[col_idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if self.titles:
                run = paragraph.add_run(self.titles[img_idx])
                run.font.bold = True
                run.font.size = Pt(self.title_fontsize)

            run = paragraph.add_run()
            if isinstance(img, str):
                run.add_picture(img, width=self.width)

            elif isinstance(img, TempOutputFile):
                run.add_picture(img.get_path(), width=self.width)
                img.delete_file()
            else:
                raise ValueError(
                    "Incorrect image source. Image sources must be a valid file path or a TempOutputFile object.")


class ObjectTable(BaseOutput):
    def __init__(self,
                 objects: List,
                 titles: Optional[List[str]] = None,
                 n_cols: int = 2,
                 width: Optional[float] = None,
                 title_fontsize: Optional[int] = 14,
                 ):
        self.objects = objects
        self.titles = titles
        self.n_cols = n_cols
        self.width = width
        self.title_fontsize = title_fontsize

    def render_pdf(self):
        pass

    def render_docx(self, document):
        self.width = Inches(self.width) if self.width is not None else self.width

        if self.titles:
            if not len(self.titles) == len(self.images):
                raise ValueError("The number of titles must match the number of objects.")

        n_rows = int(np.ceil(len(self.objects) / self.n_cols))

        table = document.add_table(rows=n_rows, cols=self.n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for obj_idx, obj in enumerate(self.objects):
            row_idx = obj_idx // self.n_cols
            col_idx = obj_idx % self.n_cols
            paragraph = table.rows[row_idx].cells[col_idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if self.titles:
                run = paragraph.add_run(self.titles[obj_idx])
                run.font.bold = True
                run.font.size = Pt(self.title_fontsize)

            if isinstance(obj, FormattedTextBlock):
                for run_obj in obj.elements:
                    run_obj.render_docx(paragraph)

            elif isinstance(obj, TempOutputFile):
                run = paragraph.add_run()
                run.add_picture(obj.get_path(), width=self.width)
                obj.delete_file()
            else:
                raise ValueError(
                    "Table objects must be either a FormattedTextBlock or a TempOutputFile object.")


class AddBreak(BaseOutput):
    def __init__(self, lines: Optional[int] = None):
        self.lines = lines if lines else 1

    def render_pdf(self):
        pass

    def render_docx(self, document):
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        self.lines = self.lines - 1
        for line in range(self.lines):
            run.add_break()


class AddPageBreak(BaseOutput):
    def __init__(self):
        pass

    def render_pdf(self):
        pass

    def render_docx(self, document):
        paragraph = document.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)
