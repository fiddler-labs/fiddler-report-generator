from .base import BaseOutput
from typing import Optional, Sequence
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches


class FormattedText(BaseOutput):
    pass


class PlainText(FormattedText):
    def __init__(self, text: str, font_color=None):
        self.text = text
        self.font_color = font_color

    def render_pdf(self):
        pass

    def render_docx(self, p):
        run = p.add_run(self.text)
        if self.font_color:
            run.font.color.rgb = self.font_color


class BoldText(FormattedText):
    def __init__(self, text: str, font_color=None):
        self.text = text
        self.font_color = font_color

    def render_pdf(self):
        pass

    def render_docx(self, p):
        run = p.add_run(self.text)
        run.font.bold = True
        if self.font_color:
            run.font.color.rgb = self.font_color


class ItalicText(FormattedText):
    def __init__(self, text: str, font_color=None):
        self.text = text
        self.font_color = font_color

    def render_pdf(self):
        pass

    def render_docx(self, p):
        run = p.add_run(self.text)
        run.font.italic = True
        if self.font_color:
            run.font.color.rgb = self.font_color


class URL(FormattedText):
    def __init__(self, text: str, url: str):
        self.text = text
        self.url = url

    def render_pdf(self):
        pass

    def render_docx(self, p):
        pass
        #it seems that python-docx does not have add_hyperlink yet. We can write something manually if needed
        #p.add_hyperlink(text=self.text, url=self.url)