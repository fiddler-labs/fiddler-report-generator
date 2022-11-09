from .base import BaseOutput
from .styles import BasicTextStyle , Alignments
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from typing import Optional

class BasicText(BaseOutput):
    def __init__(self, text:str, style:Optional[BasicTextStyle]=None):
        self.text = text
        if style:
            self.style = style
        else:
            self.style = BasicTextStyle()

    def render_pdf(self):
        pass

    def render_docx(self, document):
        #document.add_paragraph(self.text)

        paragraph = document.add_paragraph()

        if self.style.alignment == Alignments.CENTER:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = paragraph.add_run(self.text)
        run.font.bold = self.style.bold
        run.font.size = Pt(self.style.size)

