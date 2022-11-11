from .base import BaseOutput
from .styles import SimpleTextStyle, FormattedTextStyle
from .run import FormattedText
from typing import Optional, List
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches


class SimpleTextBlock(BaseOutput):
    def __init__(self, text:str, style:Optional[SimpleTextStyle]=None):
        self.text = text
        self.style = style if style else SimpleTextStyle()

    def render_pdf(self):
        pass

    def render_docx(self, document):
        paragraph = document.add_paragraph()

        if self.style.alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif self.style.alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif self.style.alignment == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif self.style.alignment == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            raise ValueError(
                'Unknown alignment parameter.'
            )

        run = paragraph.add_run(self.text)
        run.font.size = Pt(self.style.size)

        if self.style.font_style is None:
            pass
        elif self.style.font_style == 'bold':
            run.font.bold = True
        elif self.style.font_style == 'italic':
            run.font.italic = True
        else:
            raise ValueError(
                'Unknown font style.'
            )


class FormattedTextBlock(BaseOutput):
    def __init__(self, elements:List[FormattedText], style:Optional[FormattedTextStyle]=None):
        self.elements = elements
        self.style = style if style else FormattedTextStyle()

    def render_pdf(self):
        pass

    def render_docx(self, document):
        paragraph = document.add_paragraph()
        for run_obj in self.elements:
            run_obj.render_docx(paragraph)


class SimpleImage(BaseOutput):
    def __init__(self, path:str, width=None):
        self.path = path
        self.width = width

    def render_pdf(self):
        pass

    def render_docx(self, document):
        self.width= Inches(self.width) if self.width is not None else self.width
        document.add_picture(self.path, width=self.width)
